"""Export chat conversation to PDF or Word (.docx) in-memory."""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


# Special characters to normalise before writing to PDF/DOCX.
_SPECIAL_CHARS: dict[str, str] = {
    "★": "*",    # ★
    "↳": "->",   # ↳
    "·": "-",    # ·
    "—": "-",    # —
    "–": "-",    # –
    "✓": "OK",   # ✓
    "‘": "'",    # ' left single quotation mark
    "’": "'",    # ' right single quotation mark
    "“": '"',    # " left double quotation mark
    "”": '"',    # " right double quotation mark
    "R₀": "R0",  # R₀
    "₂": "2",    # ₂
}

# System CJK font candidates searched in order; first existing file wins.
_CJK_FONT_CANDIDATES = [
    # Windows
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simsun.ttc",
    # macOS
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/PingFang.ttc",
    # Linux
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttf",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
]

_RL_FONT_REGISTERED: set[str] = set()


def _replace_special(text: str) -> str:
    for char, repl in _SPECIAL_CHARS.items():
        text = text.replace(char, repl)
    return text


def _sanitize_latin(text: str) -> str:
    text = _replace_special(text)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _has_cjk(text: str) -> bool:
    return any(
        0x2E80 <= ord(c) <= 0x9FFF or 0xAC00 <= ord(c) <= 0xD7AF
        for c in text
    )


def _find_cjk_font() -> str | None:
    for path in _CJK_FONT_CANDIDATES:
        if Path(path).exists():
            return path
    return None


def _rl_register_font(font_path: str) -> str:
    """Register a TTF font with reportlab and return the font name."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "EpiChatCJK"
    if font_name not in _RL_FONT_REGISTERED:
        pdfmetrics.registerFont(TTFont(font_name, font_path))
        _RL_FONT_REGISTERED.add(font_name)
    return font_name


def to_pdf(
    messages: list[dict],
    plot_path: str | None = None,
) -> bytes:
    """Render the conversation as a PDF and return its bytes."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    all_text = " ".join(m.get("content", "") for m in messages)
    cjk_font_path = _find_cjk_font() if _has_cjk(all_text) else None

    if cjk_font_path:
        font_name = _rl_register_font(cjk_font_path)
    else:
        font_name = "Helvetica"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    title_style = ParagraphStyle(
        "title", fontName=font_name, fontSize=16, alignment=1, spaceAfter=8
    )
    label_style = ParagraphStyle(
        "label",
        fontName=font_name,
        fontSize=8,
        textColor=colors.Color(0.31, 0.31, 0.31),
        spaceAfter=2,
    )
    user_style = ParagraphStyle(
        "user", fontName=font_name, fontSize=9, spaceAfter=4
    )
    asst_style = ParagraphStyle(
        "asst",
        fontName=font_name,
        fontSize=9,
        spaceAfter=4,
        backColor=colors.Color(0.96, 0.96, 0.96),
    )

    story = [Paragraph("EpiChat Conversation", title_style), Spacer(1, 4 * mm)]

    for msg in messages:
        role = msg.get("role", "")
        raw = msg.get("content", "")
        content = _replace_special(raw)
        # Escape HTML special chars and convert newlines for Paragraph
        content = (
            content.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

        label = "You" if role == "user" else "EpiChat"
        story.append(Paragraph(label, label_style))
        story.append(
            Paragraph(content, asst_style if role == "assistant" else user_style)
        )

        msg_plot = msg.get("plot_path")
        target_plot = msg_plot or (
            plot_path
            if role == "assistant" and "complete" in raw.lower()
            else None
        )
        if target_plot and Path(target_plot).exists():
            story.append(RLImage(target_plot, width=170 * mm))

        story.append(Spacer(1, 3 * mm))

    doc.build(story)
    return buf.getvalue()


def to_docx(
    messages: list[dict],
    plot_path: str | None = None,
) -> bytes:
    """Render the conversation as a .docx file and return its bytes."""
    doc = Document()

    title = doc.add_heading("EpiChat Conversation", level=1)
    title.alignment = 1  # center

    for msg in messages:
        role = msg.get("role", "")
        content = msg.get("content", "")
        msg_plot = msg.get("plot_path")

        label_para = doc.add_paragraph()
        label_run = label_para.add_run("You" if role == "user" else "EpiChat")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(80, 80, 80)

        body_para = doc.add_paragraph(content)
        body_para.style.font.size = Pt(9)

        target_plot = msg_plot or (
            plot_path
            if role == "assistant" and "complete" in content.lower()
            else None
        )
        if target_plot and Path(target_plot).exists():
            doc.add_picture(target_plot, width=Inches(6))

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
